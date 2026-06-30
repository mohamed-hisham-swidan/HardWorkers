"""Document import pipeline.

Supports importing documents from multiple formats:
- Plain text (.txt)
- Markdown (.md)
- PDF (.pdf)
- Code files (.py, .js, .ts, .java, etc.)
- Obsidian vaults (markdown files with wikilinks)
- JSON / JSONL conversation data
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

log = logging.getLogger("hard_workers.training.document_import")


@dataclass
class ImportedDocument:
    """Normalized document after import."""

    title: str
    content: str
    source_path: str = ""
    source_type: str = "text"
    metadata: dict[str, Any] = field(default_factory=dict)
    sections: list[dict[str, Any]] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return len(self.content.split())

    @property
    def char_count(self) -> int:
        return len(self.content)


class DocumentImporter:
    """Imports documents from various file formats."""

    SUPPORTED_EXTENSIONS = {
        ".txt",
        ".md",
        ".markdown",
        ".py",
        ".js",
        ".ts",
        ".tsx",
        ".jsx",
        ".java",
        ".cpp",
        ".c",
        ".h",
        ".hpp",
        ".rs",
        ".go",
        ".rb",
        ".php",
        ".swift",
        ".kt",
        ".scala",
        ".cs",
        ".html",
        ".css",
        ".json",
        ".jsonl",
        ".yaml",
        ".yml",
        ".pdf",
        ".docx",
        ".xml",
    }

    CODE_EXTENSIONS = {
        ".py",
        ".js",
        ".ts",
        ".tsx",
        ".jsx",
        ".java",
        ".cpp",
        ".c",
        ".h",
        ".hpp",
        ".rs",
        ".go",
        ".rb",
        ".php",
        ".swift",
        ".kt",
        ".scala",
        ".cs",
        ".html",
        ".css",
    }

    # ── Public API ──────────────────────────────────────────────────────────────

    def import_file(self, path: Path | str) -> ImportedDocument:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._import_pdf(path)
        elif ext == ".docx":
            return self._import_docx(path)
        elif ext in (".json", ".jsonl"):
            return self._import_json(path)
        elif ext in (".yaml", ".yml"):
            return self._import_yaml(path)
        elif ext == ".xml":
            return self._import_xml(path)
        elif ext in (".md", ".markdown"):
            return self._import_markdown(path)
        elif ext in self.CODE_EXTENSIONS:
            return self._import_code(path)
        else:
            return self._import_text(path)

    def import_directory(
        self,
        dir_path: Path | str,
        recursive: bool = True,
        extensions: set[str] | None = None,
    ) -> list[ImportedDocument]:
        dir_path = Path(dir_path)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        extensions = extensions or self.SUPPORTED_EXTENSIONS
        docs: list[ImportedDocument] = []
        pattern = "**/*" if recursive else "*"

        for fpath in dir_path.glob(pattern):
            if fpath.is_file() and fpath.suffix.lower() in extensions:
                try:
                    doc = self.import_file(fpath)
                    docs.append(doc)
                except Exception as exc:
                    log.warning("Skipping %s: %s", fpath, exc)

        log.info("Imported %d documents from %s", len(docs), dir_path)
        return docs

    def import_obsidian_vault(
        self,
        vault_path: Path | str,
    ) -> list[ImportedDocument]:
        vault_path = Path(vault_path)
        docs = self.import_directory(vault_path, extensions={".md", ".markdown"})
        for doc in docs:
            doc.source_type = "obsidian"
            doc.metadata["wikilinks"] = self._extract_wikilinks(doc.content)
            doc.metadata["tags"] = self._extract_tags(doc.content)
        log.info("Imported Obsidian vault: %d documents from %s", len(docs), vault_path)
        return docs

    # ── Format-specific importers ───────────────────────────────────────────────

    def _import_text(self, path: Path) -> ImportedDocument:
        content = path.read_text(encoding="utf-8", errors="replace")
        return ImportedDocument(
            title=path.stem,
            content=content,
            source_path=str(path),
            source_type="text",
        )

    def _import_markdown(self, path: Path) -> ImportedDocument:
        content = path.read_text(encoding="utf-8", errors="replace")
        sections = self._split_markdown_sections(content)
        return ImportedDocument(
            title=path.stem,
            content=content,
            source_path=str(path),
            source_type="markdown",
            sections=sections,
            metadata={
                "wikilinks": self._extract_wikilinks(content),
                "tags": self._extract_tags(content),
            },
        )

    def _import_code(self, path: Path) -> ImportedDocument:
        content = path.read_text(encoding="utf-8", errors="replace")
        ext = path.suffix.lower()
        lang_map = {
            ".py": "python",
            ".js": "javascript",
            ".ts": "typescript",
            ".tsx": "typescript",
            ".jsx": "javascript",
            ".java": "java",
            ".cpp": "cpp",
            ".c": "c",
            ".h": "c",
            ".hpp": "cpp",
            ".rs": "rust",
            ".go": "golang",
            ".rb": "ruby",
            ".php": "php",
            ".swift": "swift",
            ".kt": "kotlin",
            ".scala": "scala",
            ".cs": "csharp",
            ".html": "html",
            ".css": "css",
        }
        language = lang_map.get(ext, "unknown")
        return ImportedDocument(
            title=path.stem,
            content=content,
            source_path=str(path),
            source_type="code",
            metadata={"language": language},
        )

    def _import_pdf(self, path: Path) -> ImportedDocument:
        try:
            import pypdf
        except ImportError:
            raise ImportError("pypdf is required for PDF import: pip install pypdf")
        content_parts: list[str] = []
        with open(path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content_parts.append(text)
        content = "\n\n".join(content_parts)
        return ImportedDocument(
            title=path.stem,
            content=content,
            source_path=str(path),
            source_type="pdf",
            metadata={"page_count": len(reader.pages)},
        )

    def _import_docx(self, path: Path) -> ImportedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX import: pip install python-docx")
        doc = Document(str(path))
        content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return ImportedDocument(
            title=path.stem,
            content=content,
            source_path=str(path),
            source_type="docx",
        )

    def _import_json(self, path: Path) -> ImportedDocument:
        content = path.read_text(encoding="utf-8")
        title = path.stem
        try:
            data = json.loads(content)
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
            return ImportedDocument(
                title=title,
                content=formatted,
                source_path=str(path),
                source_type="json",
            )
        except json.JSONDecodeError:
            return self._import_text(path)

    def _import_yaml(self, path: Path) -> ImportedDocument:
        try:
            import yaml
        except ImportError:
            raise ImportError("PyYAML is required for YAML import: pip install pyyaml")
        content = path.read_text(encoding="utf-8")
        try:
            data = yaml.safe_load(content)
            formatted = json.dumps(data, indent=2, ensure_ascii=False) if data else content
        except yaml.YAMLError:
            formatted = content
        return ImportedDocument(
            title=path.stem,
            content=formatted,
            source_path=str(path),
            source_type="yaml",
        )

    def _import_xml(self, path: Path) -> ImportedDocument:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 is required for XML import: pip install beautifulsoup4")
        content = path.read_text(encoding="utf-8")
        try:
            soup = BeautifulSoup(content, "lxml")
            text = soup.get_text(separator="\n", strip=True)
        except Exception:
            text = content
        return ImportedDocument(
            title=path.stem,
            content=text,
            source_path=str(path),
            source_type="xml",
        )

    # ── Content analysis helpers ────────────────────────────────────────────────

    @staticmethod
    def _extract_wikilinks(content: str) -> list[str]:
        return re.findall(r"\[\[([^\]]+)\]\]", content)

    @staticmethod
    def _extract_tags(content: str) -> list[str]:
        return list(set(re.findall(r"(?<!\w)#([a-zA-Z][a-zA-Z0-9_/-]+)", content)))

    @staticmethod
    def _split_markdown_sections(content: str) -> list[dict[str, Any]]:
        sections: list[dict[str, Any]] = []
        lines = content.split("\n")
        current_heading = "root"
        current_content: list[str] = []

        for line in lines:
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                if current_content:
                    sections.append({
                        "heading": current_heading,
                        "content": "\n".join(current_content).strip(),
                    })
                current_heading = heading_match.group(2).strip()
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            sections.append({
                "heading": current_heading,
                "content": "\n".join(current_content).strip(),
            })

        return [s for s in sections if s["content"]]
